"""Template filler for DOCX/XLSX tender artifacts."""

from __future__ import annotations
from datetime import datetime
import os
import logging

import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from openpyxl import load_workbook, Workbook
from openpyxl.cell.cell import MergedCell
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader

from app.core import helpers


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (int, float)):
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return str(value)
    return str(value)


def _normalize_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")


def build_tender_values(tender_info: Dict[str, Any]) -> Dict[str, str]:
    values = {
        "tender_id": _text(tender_info.get("tender_id")),
        "package_no": _text(tender_info.get("package_no")),
        "work_name": _text(tender_info.get("work_name") or tender_info.get("title") or tender_info.get("project_name")),
        "title": _text(tender_info.get("title") or tender_info.get("work_name")),
        "procuring_entity": _text(tender_info.get("procuring_entity") or tender_info.get("entity")),
        "division": _text(tender_info.get("division")),
        "district": _text(tender_info.get("procuring_entity_district") or tender_info.get("district")),
        "location": _text(tender_info.get("location")),
        "estimated_cost": _text(tender_info.get("estimated_cost")),
        "tender_security": _text(tender_info.get("tender_security") or tender_info.get("tender_security_amount")),
        "document_fee_bdt": _text(tender_info.get("document_fee_bdt")),
        "closing_date": _text(tender_info.get("closing_date")),
        "opening_date": _text(tender_info.get("opening_date")),
        "publication_date": _text(tender_info.get("publication_date")),
        "last_selling_date": _text(tender_info.get("last_selling_date")),
        "invitation_ref": _text(tender_info.get("invitation_ref")),
        "work_period_start": _text(tender_info.get("work_period_start")),
        "work_period_end": _text(tender_info.get("work_period_end")),
    }
    values["work_period"] = " to ".join([v for v in [values["work_period_start"], values["work_period_end"]] if v])
    return values


PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_. -]+)\s*\}\}|\[\[\s*([a-zA-Z0-9_. -]+)\s*\]\]|__([A-Z0-9_ .-]+)__")


def _extract_placeholders(text: str) -> List[str]:
    fields: List[str] = []
    for match in PLACEHOLDER_RE.finditer(text or ""):
        raw = next((group for group in match.groups() if group), "")
        key = _normalize_key(raw)
        if key:
            fields.append(key)
    return fields


def _replace_patterns(text: str, values: Dict[str, str]) -> str:
    pairs = [
        (r"(Tender/Proposal ID\s*:\s*)([^\n\r]+)", values.get("tender_id")),
        (r"(e-Tender ID\s*:\s*)([^\n\r]+)", values.get("tender_id")),
        (r"(Tender ID\s*:\s*)([^\n\r]+)", values.get("tender_id")),
        (r"(Package No\.?\s*:\s*)([^\n\r\)]*)", values.get("package_no")),
        (r"(Package\s*:\s*)([^\n\r\)]*)", values.get("package_no")),
        (r"(Name of Work\s*:\s*)([^\n\r]+)", values.get("work_name") or values.get("title")),
        (r"(Work\s*Name\s*:\s*)([^\n\r]+)", values.get("work_name") or values.get("title")),
        (r"(Procuring Entity(?: Name)?\s*:\s*)([^\n\r]+)", values.get("procuring_entity")),
        (r"(Division\s*:\s*)([^\n\r]+)", values.get("division")),
        (r"(District\s*:\s*)([^\n\r]+)", values.get("district")),
        (r"(Tender Security(?: Amount)?\s*:\s*)([^\n\r]+)", values.get("tender_security")),
        (r"(Security Amount\s*:\s*)([^\n\r]+)", values.get("tender_security")),
        (r"(Line of Credit(?: Amount)?\s*:\s*)([^\n\r]+)", values.get("credit_line_amount")),
    ]
    for pattern, replacement in pairs:
        if not replacement:
            continue
        text = re.sub(pattern, lambda m: f"{m.group(1)}{replacement}", text, flags=re.IGNORECASE)

    for key, replacement in values.items():
        if not replacement:
            continue
        text = text.replace(f"{{{{{key}}}}}", replacement)
        text = re.sub(r"\{\{\s*" + re.escape(key) + r"\s*\}\}", replacement, text, flags=re.IGNORECASE)
        text = text.replace(f"[[{key}]]", replacement)
        text = re.sub(r"\[\[\s*" + re.escape(key) + r"\s*\]\]", replacement, text, flags=re.IGNORECASE)
        text = text.replace(f"__{key.upper()}__", replacement)
    return text


def _iter_doc_text(doc: Document) -> Iterable[str]:
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph.text
        for paragraph in section.footer.paragraphs:
            yield paragraph.text


def _write_validation_report(output_path: Path, report: Dict[str, Any]) -> str:
    report_path = output_path.with_suffix(".validation.json")
    helpers.write_json(report_path, report)
    return str(report_path)


def fill_docx_template(template_path: str, output_path: str, values: Dict[str, Any], extra: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    mapping = build_tender_values(values)
    if extra:
        for key, value in extra.items():
            mapping[_normalize_key(key)] = _text(value)

    doc = Document(str(template_path))
    detected = sorted(set(field for text in _iter_doc_text(doc) for field in _extract_placeholders(text)))
    before_text = "\n".join(_iter_doc_text(doc))

    def update_paragraph(paragraph):
        if not paragraph.text:
            return
        new_text = _replace_patterns(paragraph.text, mapping)
        if new_text != paragraph.text:
            paragraph.text = new_text

    for paragraph in doc.paragraphs:
        update_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            update_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            update_paragraph(paragraph)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    after_doc = Document(str(out))
    remaining = sorted(set(field for text in _iter_doc_text(after_doc) for field in _extract_placeholders(text)))
    filled = sorted(field for field in detected if field not in remaining and mapping.get(field))
    missing = sorted(field for field in detected if field in remaining or not mapping.get(field))
    pattern_filled = sorted(key for key, value in mapping.items() if value and str(value) in "\n".join(_iter_doc_text(after_doc)) and str(value) not in before_text)
    report = {
        "template": str(template_path),
        "output_path": str(out),
        "detected_placeholders": detected,
        "filled_placeholders": filled,
        "missing_placeholders": missing,
        "pattern_filled_fields": pattern_filled,
        "all_fields": mapping,
    }
    report_path = _write_validation_report(out, report)
    return {"output_path": str(out), "fields": mapping, "report": report, "report_path": report_path}


def _set_if_present(ws, cell_ref: str, value: Any):
    if value is None:
        return
    cell = ws[cell_ref]
    if isinstance(cell, MergedCell):
        for merged_range in ws.merged_cells.ranges:
            if cell.coordinate in merged_range:
                ws.cell(merged_range.min_row, merged_range.min_col).value = value
                return
    cell.value = value


def fill_workbook_template(
    template_path: str,
    output_path: str,
    tender_info: Dict[str, Any],
    boq_items: Optional[List[Dict[str, Any]]] = None,
    sor_service: Any = None,
    agency: str = "BWDB",
    zone: Optional[str] = None,
) -> Dict[str, Any]:
    wb = load_workbook(str(template_path))
    values = build_tender_values(tender_info)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    updated = _replace_patterns(cell.value, values)
                    if updated != cell.value:
                        cell.value = updated

        title = ws.title.lower()
        if title == "quot":
            _populate_boq_sheet(ws, values, boq_items or [], sor_service, agency, zone)
        elif _looks_like_work_plan(ws, title):
            _populate_work_plan_sheet(ws, values)
        elif title == "sum_quoted":
            _set_if_present(ws, "D2", values.get("tender_id"))

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out))
    return {"output_path": str(out), "fields": values}


def _populate_boq_sheet(ws, values: Dict[str, str], boq_items: List[Dict[str, Any]], sor_service: Any, agency: str, zone: Optional[str]):
    _set_if_present(ws, "B2", f"Tender ID:{values.get('tender_id')}")
    _set_if_present(ws, "D2", values.get("procuring_entity"))
    _set_if_present(ws, "B6", values.get("package_no"))
    _set_if_present(ws, "D6", values.get("work_name") or values.get("title"))
    _set_if_present(ws, "E6", values.get("district"))
    _set_if_present(ws, "F6", values.get("estimated_cost"))
    _set_if_present(ws, "H6", values.get("work_period") or values.get("closing_date"))

    by_code = {}
    for item in boq_items:
        key = _normalize_key(_text(item.get("code")))
        if key:
            by_code[key] = item

    start_row = 10
    for row_idx in range(start_row, ws.max_row + 1):
        code = _text(ws[f"C{row_idx}"].value).strip()
        item = None
        if code:
            item = by_code.get(_normalize_key(code))
        if not item and (row_idx - start_row) < len(boq_items):
            item = boq_items[row_idx - start_row]
        if not item:
            continue

        _set_if_present(ws, f"B{row_idx}", item.get("item_no") or row_idx - start_row + 1)
        _set_if_present(ws, f"C{row_idx}", item.get("code"))
        _set_if_present(ws, f"D{row_idx}", item.get("description"))
        _set_if_present(ws, f"E{row_idx}", item.get("quantity"))
        _set_if_present(ws, f"F{row_idx}", item.get("unit"))
        rate, record = (None, None)
        if sor_service:
            rate, record = sor_service.find_rate(_text(item.get("code")), _text(item.get("description")), agency, zone)
        if rate is not None:
            ws[f"G{row_idx}"] = rate
            ws[f"H{row_idx}"] = f"=G{row_idx}*E{row_idx}"


def _populate_work_plan_sheet(ws, values: Dict[str, str]):
    _set_if_present(ws, "B3", f"Tender ID:   {values.get('tender_id')}")
    name = values.get("work_name") or values.get("title")
    if name:
        _set_if_present(ws, "B4", f'Name of Work:“{name}”')
    if values.get("package_no"):
        _set_if_present(ws, "O3", f"Package: {values.get('package_no')}")


def _looks_like_work_plan(ws, title: str) -> bool:
    if "work" in title and "plan" in title:
        return True
    probe_cells = ["B2", "B3", "B4", "E3", "O3", "B5", "E5"]
    probe = " ".join(_text(ws[cell].value) for cell in probe_cells)
    probe = probe.lower()
    return any(marker in probe for marker in ["work schedule", "tender id", "name of work", "package:"])


def create_basic_work_plan(output_path: str, tender_info: Dict[str, Any]) -> str:
    wb = Workbook()
    ws = wb.active
    ws.title = str(tender_info.get("tender_id") or "Work Plan")
    values = build_tender_values(tender_info)
    ws["B3"] = f"Tender ID:   {values.get('tender_id')}"
    ws["B4"] = f'Name of Work:“{values.get("work_name") or values.get("title")}”'
    ws["O3"] = f"Package: {values.get('package_no')}"
    ws["E3"] = "WORK SCHEDULE"
    wb.save(output_path)
    return output_path


def create_pdf_text_docx(pdf_path: str, output_path: str, title: str = "") -> str:
    doc = Document()
    heading = title or Path(pdf_path).stem
    doc.add_heading(heading, level=1)
    doc.add_paragraph(f"Source PDF: {Path(pdf_path).name}")

    with open(pdf_path, "rb") as handle:
        reader = PdfReader(handle)
        for index, page in enumerate(reader.pages, start=1):
            if index > 1:
                doc.add_page_break()
            doc.add_heading(f"Page {index}", level=2)
            text = page.extract_text() or ""
            for raw_line in text.splitlines():
                line = re.sub(r"\s+", " ", raw_line).strip()
                if not line:
                    continue
                paragraph = doc.add_paragraph(line)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ─────────────────────────────────────────────────────────────────────
# Tender Document Template Filler (for uploaded DOCX templates)
# Uses actual tender data to replace old values in templates
# ─────────────────────────────────────────────────────────────────────

logger = logging.getLogger(__name__)


TEMPLATE_FIELD_MAPS: Dict[str, List[tuple]] = {
    "JV_DEED.docx": [
        (r"556156|556145", "$tender_id"),
        (r"Nao\-W\-5|BSL/MHDJ/WD\-11", "$package_no"),
        (r"EE/NO&MD/Da\-2/[\d\s\-]+", "$invitation_ref_short"),
        (r"Bank Revetment work along the Left bank of the Atrai River[^.]+", "$project_description"),
        (r"River bank protection work of Tetulia river[^.]+", "$project_description"),
        (r"5th April,2021", "$agreement_date"),
        (r"06\-04\-2021", "$agreement_date"),
        (r"Mahmud Tower[^,]+[^)]+\)", "$bidder_address"),
        (r"02\-9564202, 9563940", "$bidder_phone"),
        (r"hbl\.engr@gmail\.com", "$bidder_email"),
        (r"Md\. Shamimur Rahman", "$jv_partner2_name"),
        (r"Kalibari Road, Sirajganj[^.]+", "$jv_partner2_address"),
        (r"60% \(Sixty percent\)", "$jv_share_1"),
        (r"40% \(Forty percent\)", "$jv_share_2"),
    ],
    "Credit_Line.docx": [
        (r"556156", "$tender_id"),
        (r"Nao\-W\-5", "$package_no"),
        (r"EE/NO&MD/Da\-2/[\d\s\-]+", "$invitation_ref_short"),
        (r"Hassan & Brothers", "$bidder_name"),
        (r"NRB Commercial Bank Ltd[^,]*", "$bank_name_full"),
        (r"Principal Branch", "$bank_branch"),
        (r"114, Motijheel C/A", "$bank_address"),
    ],
    "BG_Tender_Security.docx": [
        (r"556156", "$tender_id"),
        (r"Nao\-W\-5", "$package_no"),
        (r"EE/NO&MD/Da\-2/[\d\s\-]+", "$invitation_ref_short"),
        (r"13,00,000[^)]*", "$tender_security_text"),
        (r"Taka Thirteen Lac", "$tender_security_words"),
        (r"Hassan & Brothers", "$bidder_name"),
        (r"HB\-SR\-MT JV[^,]*", "$jv_name"),
        (r"NRB Commercial Bank Ltd[^,]*", "$bank_name_full"),
    ],
    "BG_Application.docx": [
        (r"HB/4564", "$memo_no"),
        (r"06\-04\-2021", "$application_date"),
        (r"NRB Commercial Bank Limited", "$bank_name_full"),
        (r"Principal Branch", "$bank_branch"),
        (r"114, Motijheel C/A, Dhaka", "$bank_address"),
        (r"Hassan & Brothers", "$bidder_name"),
    ],
    "Equipment_Declaration.docx": [
        (r"River bank protection work of Tetulia river[^\"]*", "$project_description"),
        (r"556156", "$tender_id"),
    ],
    "Manpower_Declaration.docx": [
        (r"River bank protection work of Tetulia river[^\"]*", "$project_description"),
        (r"556156", "$tender_id"),
    ],
    "Methodology.docx": [
        (r"River bank protection work of Tetulia river[^.]+", "$project_description"),
        (r"1264860", "$tender_id"),
        (r"BSL/MHDJ/WD\-11", "$package_no"),
        (r"Barishal WD Division, BWDB", "$procuring_entity"),
    ],
    "JV_POWER.docx": [
        (r"HB\-SR JV\.", "$jv_name"),
        (r"Hassan & Brothers", "$bidder_name"),
        (r"Mahmudul Hassan", "$bidder_representative"),
        (r"Md\. Shamimur Rahman", "$jv_partner2_name"),
        (r"Kalibari Road, Sirajganj", "$jv_partner2_address"),
    ],
}


def _get_field_value(key: str, data: Dict[str, Any]) -> str:
    """Get field value from data dict. Return placeholder if not found."""
    if not key.startswith("$"):
        return key
    field = key[1:]
    val = data.get(field, "")
    return str(val) if val is not None else ""


def _replace_in_docx_element(element, field_map: List[tuple], data: Dict[str, Any]) -> bool:
    """Replace text in a docx element (paragraph, run, cell) using field map."""
    modified = False
    for pattern, replacement_key in field_map:
        replacement = _get_field_value(replacement_key, data)
        if not replacement:
            continue
        try:
            # Handle both paragraph and run level
            if hasattr(element, 'text') and re.search(pattern, element.text or ''):
                new_text = re.sub(pattern, replacement, element.text or '')
                element.text = new_text
                modified = True
        except Exception:
            continue
    return modified


def fill_docx_from_template(
    template_path: str,
    output_path: str,
    data: Dict[str, Any],
    template_name: str = "",
) -> str:
    """Fill a DOCX template with tender data.
    
    Uses TEMPLATE_FIELD_MAPS to replace old tender values with new data.
    
    Args:
        template_path: Path to DOCX template
        output_path: Output path for generated document
        data: Dict of replacement values
        template_name: Template identifier for field map lookup
    
    Returns:
        Path to generated document, or empty string if template not found
    """
    if not os.path.exists(template_path):
        logger.warning(f"Template not found: {template_path}")
        return ""
    
    fname = template_name or os.path.basename(template_path)
    field_map = TEMPLATE_FIELD_MAPS.get(fname, [])
    
    doc = Document(template_path)
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        _replace_in_docx_element(p, field_map, data)
        for run in p.runs:
            _replace_in_docx_element(run, field_map, data)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_docx_element(p, field_map, data)
                    for run in p.runs:
                        _replace_in_docx_element(run, field_map, data)
    
    # Replace in headers/footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_docx_element(p, field_map, data)
        for p in section.footer.paragraphs:
            _replace_in_docx_element(p, field_map, data)
    
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info(f"Generated: {output_path}")
    return output_path


def generate_all_tender_docs(
    data: Dict[str, Any],
    templates_dir: str,
    output_dir: str = "./output",
) -> Dict[str, str]:
    """Generate all tender documents from templates.
    
    Args:
        data: Field values dict
        templates_dir: Directory with template DOCX files
        output_dir: Output directory
    
    Returns:
        Dict of filename -> status
    """
    os.makedirs(output_dir, exist_ok=True)
    results = {}
    
    template_pairs = [
        ("JV_DEED.docx", "JV_DEED.docx"),
        ("Credit_Line.docx", "1_Credit_Line.docx"),
        ("BG_Tender_Security.docx", "2_BG_Tender_Security.docx"),
        ("BG_Application.docx", "3_BG_Application.docx"),
        ("Equipment_Declaration.docx", "4_Equipment_Declaration.docx"),
        ("Manpower_Declaration.docx", "5_Manpower_Declaration.docx"),
        ("Methodology.docx", "6_Methodology.docx"),
        ("JV_POWER.docx", "7_JV_POWER_OF_ATTORNEY.docx"),
    ]
    
    for tpl_name, out_name in template_pairs:
        tpl_path = os.path.join(templates_dir, tpl_name)
        out_path = os.path.join(output_dir, out_name)
        try:
            result = fill_docx_from_template(tpl_path, out_path, data, tpl_name)
            results[out_name] = "✅" if result else "⚠️ Template missing"
        except Exception as e:
            logger.error(f"Failed {out_name}: {e}")
            results[out_name] = f"❌ {str(e)[:50]}"
    
    return results


def build_tender_template_data(
    tender_data: Any = None,
    bidder_name: str = "",
    bidder_address: str = "",
    bidder_phone: str = "",
    bidder_email: str = "",
    bank_name: str = "",
    bank_branch: str = "",
    bank_address: str = "",
    jv_name: str = "",
    jv_partner2_name: str = "",
    jv_partner2_address: str = "",
    bidder_representative: str = "",
) -> Dict[str, Any]:
    """Build comprehensive data dict for template filling."""
    td = {}
    if tender_data is not None:
        if hasattr(tender_data, '__dict__'):
            td = tender_data.__dict__
        elif hasattr(tender_data, '_asdict'):
            td = tender_data._asdict()
        elif isinstance(tender_data, dict):
            td = tender_data
    
    data = {
        "tender_id": str(td.get("tender_id", td.get("id", "1264860"))),
        "package_no": str(td.get("package_no", "BSL/MHDJ/WD-11")),
        "invitation_ref": str(td.get("invitation_ref", "")),
        "invitation_ref_short": str(td.get("invitation_ref", "42.01.0600.000.112.07.0001.25-2642, Date-27/04/26"))[:60],
        "project_description": str(td.get("package_description", td.get("brief", "River Bank Protection Work"))),
        "procuring_entity": str(td.get("procuring_entity", "BWDB")),
        "organization": str(td.get("organization", "")),
        "estimated_value_text": str(td.get("estimated_value_text", "")),
        "tender_security_text": str(td.get("tender_security_text", "Tk. 61,00,000 (Sixty One Lakh)")),
        "tender_security_words": "Sixty One Lakh",
        "agreement_date": datetime.now().strftime("%d %B, %Y"),
        "application_date": datetime.now().strftime("%d-%m-%Y"),
        "bidder_name": bidder_name or "M/S. Hassan & Brothers",
        "bidder_address": bidder_address or "Mahmud Tower, (9th Floor) 19, Siddique Bazar, Dhaka",
        "bidder_phone": bidder_phone or "02-9564202, 9563940",
        "bidder_email": bidder_email or "hbl.engr@gmail.com",
        "bidder_representative": bidder_representative or "Mahmudul Hassan",
        "bank_name_full": bank_name or "NRB Commercial Bank Ltd.",
        "bank_branch": bank_branch or "Principal Branch",
        "bank_address": bank_address or "114, Motijheel C/A, Dhaka",
        "jv_name": jv_name or "HB-SR JV.",
        "jv_partner2_name": jv_partner2_name or "Md. Shamimur Rahman",
        "jv_partner2_address": jv_partner2_address or "Kalibari Road, Sirajganj, Bangladesh",
        "jv_share_1": "60% (Sixty percent)",
        "jv_share_2": "40% (Forty percent)",
        "memo_no": f"HB/{datetime.now().strftime('%Y%m%d')}",
    }
    return data
