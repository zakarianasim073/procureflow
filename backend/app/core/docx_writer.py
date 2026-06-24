"""BOQ DOCX Report Writer - Generates professional Word document analysis"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime


def set_cell_shading(cell, hex_color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1e40af')

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT
            text = str(value) if value is not None else ''
            run = p.add_run(text)
            run.font.size = Pt(9)

            # Alternate row shading
            if r % 2 == 1:
                set_cell_shading(cell, 'f3f4f6')

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def write_boq_docx(
    filepath: str,
    tender_info: dict,
    comparison_rows: list,
    summary_rows: list,
    flagged_rows: list,
    financial_rows: list,
):
    """Generate professional Word document with full BOQ analysis."""

    doc = Document()

    # ── Styles ───────────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 1: Cover / Header
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BOQ vs SOR RATE ANALYSIS REPORT')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Tender ID: {tender_info.get('tender_id', 'N/A')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = tender_info.get('title', '')
    if len(title) > 200:
        title = title[:200] + '...'
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    # Tender Information Table
    info_fields = [
        ('Procuring Entity', tender_info.get('entity')),
        ('Location', tender_info.get('location')),
        ('Agency / SOR', tender_info.get('sor_agency')),
        ('Zone', tender_info.get('zone')),
        ('Invitation Ref.', tender_info.get('invitation_ref')),
        ('Tender ID', tender_info.get('tender_id')),
        ('Tender Security', tender_info.get('tender_security')),
        ('Analysis Date', datetime.now().strftime('%d-%b-%Y %H:%M')),
    ]
    info_table = doc.add_table(rows=len(info_fields), cols=2)
    info_table.style = 'Light Shading Accent 1'
    for i, (label, value) in enumerate(info_fields):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value or 'N/A')
        for cell in info_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True if info_table.rows[i].cells[0].paragraphs[0].runs else False

    doc.add_paragraph()  # spacer

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 2: Key Metrics Summary
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run('KEY METRICS')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    total_sor = tender_info.get('total_sor', 0)
    total_quoted = tender_info.get('total_quoted', 0)
    saving = tender_info.get('saving', 0)
    discount_pct = tender_info.get('discount_pct', 0)

    metrics_rows = [
        ('Total SOR Amount', f'BDT {total_sor:,.2f}'),
        ('Total Quoted Amount', f'BDT {total_quoted:,.2f}'),
        ('Net Saving' if saving >= 0 else 'Net Overrun', f'BDT {saving:,.2f}'),
        ('Overall Discount', f'{discount_pct:.2%}'),
        ('Flagged Items', str(len(flagged_rows))),
        ('Items Matched', str(sum(1 for r in comparison_rows if r.get('flag') == 'AT SOR'))),
        ('Total Items', str(len(comparison_rows))),
    ]
    add_styled_table(doc, ['Metric', 'Value'], metrics_rows, col_widths=[10, 8])

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 3: Work Type Summary
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run('WORK TYPE SUMMARY')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    if summary_rows:
        headers = ['Work Type', 'Items', 'SOR Amount (BDT)', 'Quoted Amount (BDT)', 'Saving (BDT)', 'Discount %']
        data = []
        for r in summary_rows:
            data.append([
                r.get('work_type', ''),
                str(r.get('items', 0)),
                f"{r.get('sor_amount', 0):,.2f}",
                f"{r.get('quoted_amount', 0):,.2f}",
                f"{r.get('saving', 0):,.2f}",
                f"{r.get('discount_pct', 0)*100:.1f}%",
            ])
        add_styled_table(doc, headers, data, col_widths=[3, 1.5, 3, 3, 3, 2])
    else:
        doc.add_paragraph('No work type summary available.')

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 4: BOQ Item Comparison (Detailed)
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run('DETAILED BOQ ITEM COMPARISON')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    if comparison_rows:
        headers = ['#', 'Code', 'Description', 'Unit', 'Qty', 'Rate (BDT)', 'SOR Rate (BDT)', 'Diff %', 'Status']
        data = []
        for item in comparison_rows[:50]:  # Limit to first 50 items to avoid huge doc
            flag = item.get('flag', '')
            pct = item.get('pct_diff')
            data.append([
                str(item.get('item_no', '')),
                item.get('code', ''),
                (item.get('desc', '') or '')[:60],
                item.get('unit', ''),
                str(item.get('qty', '')),
                f"{item.get('rate', 0) or 0:,.2f}",
                f"{item.get('sor_rate', 0) or 0:,.2f}",
                f"{pct:+.1f}%" if pct is not None else 'N/A',
                flag,
            ])
        add_styled_table(doc, headers, data, col_widths=[1, 2.5, 5, 1.5, 1.5, 2, 2, 1.5, 2])

        if len(comparison_rows) > 50:
            doc.add_paragraph(f'Note: Showing first 50 of {len(comparison_rows)} items.')
    else:
        doc.add_paragraph('No comparison data available.')

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 5: Flagged Items (DISCREPANCIES)
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run('FLAGGED ITEMS — REQUIRES ATTENTION')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    if flagged_rows:
        headers = ['#', 'Code', 'Description', 'Qty', 'Rate', 'SOR Rate', 'Diff %', 'Flag']
        data = []
        for item in flagged_rows:
            pct = item.get('pct_diff')
            data.append([
                str(item.get('item_no', '')),
                item.get('code', ''),
                (item.get('desc', '') or '')[:50],
                str(item.get('qty', '')),
                f"{item.get('rate', 0) or 0:,.2f}",
                f"{item.get('sor_rate', 0) or 0:,.2f}",
                f"{pct:+.1f}%" if pct is not None else 'N/A',
                item.get('flag', ''),
            ])
        add_styled_table(doc, headers, data, col_widths=[1, 2.5, 4.5, 1.5, 2, 2, 1.5, 2])
    else:
        doc.add_paragraph('No flagged items — all rates match SOR within acceptable range.')

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 6: Financial Check
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run('FINANCIAL ELIGIBILITY CHECK')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    if financial_rows:
        headers = ['Criterion', 'Required', 'Offered', 'Status']
        data = []
        for item in financial_rows:
            data.append([
                item.get('criterion', ''),
                item.get('required', ''),
                item.get('offered', ''),
                item.get('status', ''),
            ])
        add_styled_table(doc, headers, data, col_widths=[5, 4, 4, 3])
    else:
        doc.add_paragraph('No financial eligibility data available.')

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 7: Footer / Disclaimers
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— End of Report —')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated by Procurement Flow Specialist BD v2.0 on {datetime.now().strftime("%d-%b-%Y %H:%M")}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SOR rates are for reference only. Verify all rates with official schedule before bidding.')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Save
    doc.save(str(filepath))
    print(f"  DOCX report saved: {filepath}")
    return filepath
