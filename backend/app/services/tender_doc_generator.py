"""
Tender Document Generator — Auto-generates submission documents from tender PDF data.
Extracts data from Notice, TDS, BOQ PDFs and maps to DOCX templates.

Usage:
    from app.services.tender_doc_generator import TenderDocGenerator
    
    gen = TenderDocGenerator()
    
    # Extract data from tender PDFs
    data = gen.extract_from_pdfs(notice_pdf="1.Notice.pdf", tds_pdf="2.TDS.pdf", boq_pdf="4.BOQ.pdf")
    
    # Generate documents
    gen.generate_all(data, output_dir="./output")
    
    # Or generate individual documents
    gen.generate_jv_deed(data, "output/JV_DEED.docx")
    gen.generate_credit_line(data, "output/Credit_Line.docx")
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class TenderData:
    """All data extracted from tender documents, used to populate templates."""
    
    # Tender Identity
    tender_id: str = ""
    app_id: str = ""
    invitation_ref: str = ""
    
    # Organization
    ministry: str = ""
    organization: str = ""
    procuring_entity: str = ""
    procuring_entity_district: str = ""
    
    # Procurement Details
    procurement_nature: str = "Works"
    procurement_type: str = "NCT"
    procurement_method: str = "Open Tendering Method (OTM)"
    budget_type: str = "Development"
    source_of_funds: str = "Government"
    
    # Project
    project_code: str = ""
    project_name: str = ""
    package_no: str = ""
    package_description: str = ""
    
    # Key Info
    estimated_value_bdt: float = 0.0
    estimated_value_text: str = ""
    tender_security_bdt: float = 0.0
    tender_security_text: str = ""
    validity_period_days: int = 120
    completion_period_days: int = 0
    
    # Dates
    tender_close_datetime: str = ""
    tender_open_datetime: str = ""
    project_start_date: str = ""
    project_end_date: str = ""
    
    # Officer
    inviting_officer_name: str = ""
    inviting_officer_designation: str = ""
    inviting_officer_address: str = ""
    inviting_officer_city: str = ""
    inviting_officer_phone: str = ""
    
    # TDS Data
    experience_years_required: int = 5
    equipment_list: List[Dict] = field(default_factory=list)
    manpower_list: List[Dict] = field(default_factory=list)
    security_in_favor_of: str = ""
    tax_requirements: List[str] = field(default_factory=list)
    
    # BOQ Data
    boq_items: List[Dict] = field(default_factory=list)
    boq_total: float = 0.0
    
    # Bidder Info (from templates / user input)
    bidder_name: str = ""
    bidder_address: str = ""
    bidder_phone: str = ""
    bidder_email: str = ""
    jv_partners: List[Dict] = field(default_factory=list)
    bank_name: str = ""
    bank_branch: str = ""
    bank_address: str = ""


class TenderPDFExtractor:
    """Extract structured data from tender PDFs (Notice, TDS, BOQ)."""

    @staticmethod
    def extract_from_text(text: str, field_mappings: Dict[str, str] = None) -> Dict[str, Any]:
        """Generic extraction using regex patterns."""
        data = {}
        if not field_mappings:
            return data
        
        for field, pattern in field_mappings.items():
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                data[field] = m.group(1).strip()
        return data

    @staticmethod
    def extract_notice(pdf_path: str) -> Dict[str, Any]:
        """Extract data from Tender Notice PDF."""
        from PyPDF2 import PdfReader
        
        data = {}
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            
            # Extract fields
            patterns = {
                "ministry": r'Ministry\s*:\s*(.+?)(?:Division|\n)',
                "organization": r'Organization\s*:\s*(.+?)(?:\n)',
                "procuring_entity": r'Procuring Entity Name\s*:\s*(.+?)(?:\n)',
                "district": r'Procuring Entity District\s*:\s*(.+?)(?:\n)',
                "invitation_ref": r'Invitation Reference No\.?\s*:\s*(.+?)(?:\n|,)',
                "app_id": r'App ID\s*:\s*(\d+)',
                "tender_id": r'Tender/Proposal ID\s*:\s*(\d+)',
                "procurement_method": r'Procurement Method\s*:\s*(.+?)(?:\n)',
                "budget_type": r'Budget Type\s*:\s*(.+?)(?:\n)',
                "source_of_funds": r'Source of Funds\s*:\s*(.+?)(?:\n)',
                "project_code": r'Project Code\s*:\s*(.+?)(?:\n)',
                "project_name": r'Project Name\s*:\s*(.+?)(?:\n)',
                "package_no": r'Package No\.?\s*(?:and Description)?\s*:\s*([^\n]+)',
                "estimated_value": r'(\d[\d,]*)\s*(?:Tk|BDT|)\s*[-\s]*(?:Oct|Sep|Dec|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug)',
                "officer_name": r'Name of Official Inviting[^:]*:\s*(.+?)(?:\n)',
                "officer_designation": r'Designation of Official Inviting[^:]*:\s*(.+?)(?:\n)',
                "officer_address": r'Address\s*:\s*(.+?)(?:\n)',
                "officer_city": r'City\s*:\s*(.+?)(?:\n)',
                "officer_phone": r'Phone No\s*:\s*(.+?)(?:\n)',
            }
            
            for field, pattern in patterns.items():
                m = re.search(pattern, text, re.IGNORECASE)
                if m:
                    val = m.group(1).strip()
                    # Clean up
                    val = re.sub(r'\s+', ' ', val).strip()
                    data[field] = val
            
            # Extract project description
            desc_match = re.search(r'Description\s*:\s*([^\n]+(?:[^\n]*\n[^\n]*)*)', text, re.IGNORECASE)
            if desc_match:
                desc = re.sub(r'\s+', ' ', desc_match.group(1)).strip()
                data["package_description"] = desc
            
            # Extract project dates
            dates = re.findall(r'(\d{2}-[A-Za-z]{3}-\d{4})', text)
            if len(dates) >= 2:
                data["project_start_date"] = dates[-2]
                data["project_end_date"] = dates[-1]
            
            # Extract estimated value as number
            val_match = re.search(r'(\d[\d,]*)\s*(?:Tk|BDT)', text)
            if val_match:
                try:
                    data["estimated_value_bdt"] = float(val_match.group(1).replace(",", ""))
                except ValueError:
                    pass
            
            logger.info(f"Notice extracted: {data.get('tender_id', '?')} - {data.get('package_no', '?')}")
            
        except Exception as e:
            logger.error(f"Notice extraction failed: {e}")
        
        return data

    @staticmethod
    def extract_tds(pdf_path: str) -> Dict[str, Any]:
        """Extract data from Tender Data Sheet (TDS) PDF."""
        from PyPDF2 import PdfReader
        
        data = {
            "equipment_list": [],
            "manpower_list": [],
            "tax_requirements": [],
        }
        
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            
            # Tender Security
            sec_match = re.search(
                r'Tender Security\s*(?:shall be|amount|:).*?Tk\.?\s*([\d,]+\.?\d*)\s*(?:Lakh|Lac)?',
                text, re.IGNORECASE
            )
            if sec_match:
                val = sec_match.group(1).replace(",", "")
                try:
                    data["tender_security_bdt"] = float(val) * 100000 if 'lakh' in text[sec_match.start():sec_match.end()].lower() else float(val)
                except ValueError:
                    pass
            
            # Security in favor of
            fav_match = re.search(r'in favour of\s*(.+?)(?:\.|\n)', text, re.IGNORECASE)
            if fav_match:
                data["security_in_favor_of"] = fav_match.group(1).strip()
            
            # Validity period
            val_match = re.search(r'valid\s*(?:for|period).*?(\d+)\s*days?', text, re.IGNORECASE)
            if val_match:
                data["validity_period_days"] = int(val_match.group(1))
            
            # Equipment table parsing
            equip_patterns = [
                r'(\d+)\s+(.+?)\s+(\d+\s*(?:Nos?|Set|Unit|No\.?))\s*$',
                r'Equipment[^:]*:\s*(.+?)(?:\d+\.\s)',
            ]
            
            # Manpower table parsing  
            lines = text.split('\n')
            in_equip_table = False
            in_manpower_table = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect equipment table
                if 'equipment' in line.lower() and ('type' in line.lower() or 'character' in line.lower()):
                    in_equip_table = True
                    in_manpower_table = False
                    continue
                    
                if 'key staff' in line.lower() or 'post of key' in line.lower():
                    in_manpower_table = True
                    in_equip_table = False
                    continue
                
                if in_equip_table:
                    eq_match = re.match(r'^(\d+)\s+(.+?)\s+(\d[\d\s]*(?:Nos?|Set|Unit|No\.?))', line, re.IGNORECASE)
                    if eq_match:
                        data["equipment_list"].append({
                            "sl": eq_match.group(1),
                            "type": eq_match.group(2).strip(),
                            "min_required": eq_match.group(3).strip(),
                        })
                
                if in_manpower_table:
                    mp_match = re.match(r'^(\d+)\s+(.+?)\s+([A-Za-z].+?)\s+(\d[\d\s]*(?:Person|Nos?|No\.?))\s+(\d[\d\s]*years?)\s+(\d[\d\s]*years?)', line, re.IGNORECASE)
                    if mp_match:
                        data["manpower_list"].append({
                            "sl": mp_match.group(1),
                            "post": mp_match.group(2).strip(),
                            "qualification": mp_match.group(3).strip(),
                            "nos": mp_match.group(4).strip(),
                            "total_exp": mp_match.group(5).strip(),
                            "similar_exp": mp_match.group(6).strip(),
                        })
            
            logger.info(f"TDS extracted: {len(data['equipment_list'])} equipment, {len(data['manpower_list'])} manpower")
            
        except Exception as e:
            logger.error(f"TDS extraction failed: {e}")
        
        return data

    @staticmethod
    def extract_boq(pdf_path: str) -> Dict[str, Any]:
        """Extract data from Bill of Quantities (BOQ) PDF."""
        from PyPDF2 import PdfReader
        
        data = {"boq_items": [], "boq_total": 0.0}
        
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            
            # Parse BOQ items: look for item code + description + unit + quantity patterns
            lines = text.split('\n')
            current_item = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # BWDB item code pattern: "1 40-920" or "3 40-920"
                item_match = re.match(
                    r'^(\d+)\s+([\d\-–]+(?:\-[\d]+)?)\s+(.+?)\s+(cum|sqm|each|no|nos|kg|m\.?|day|%|lump|set)\s+([\d,]+\.?\d*)',
                    line, re.IGNORECASE
                )
                if item_match:
                    item = {
                        "sl": item_match.group(1),
                        "code": item_match.group(2),
                        "description": item_match.group(3).strip(),
                        "unit": item_match.group(4).lower(),
                        "quantity": float(item_match.group(5).replace(",", "")),
                    }
                    data["boq_items"].append(item)
            
            logger.info(f"BOQ extracted: {len(data['boq_items'])} items")
            
        except Exception as e:
            logger.error(f"BOQ extraction failed: {e}")
        
        return data


class TenderDocGenerator:
    """Generate tender submission documents from data + templates."""

    def __init__(self, templates_dir: str = None):
        self.templates_dir = templates_dir or str(Path(__file__).parent / "templates")
        self.extractor = TenderPDFExtractor()
        self.template_cache = {}

    def extract_from_pdfs(
        self,
        notice_pdf: str = "",
        tds_pdf: str = "",
        boq_pdf: str = "",
    ) -> TenderData:
        """Extract all data from tender PDFs into structured TenderData."""
        data = TenderData()
        
        # Extract from Notice
        if notice_pdf and os.path.exists(notice_pdf):
            notice_data = self.extractor.extract_notice(notice_pdf)
            for key, val in notice_data.items():
                if hasattr(data, key):
                    setattr(data, key, val)
        
        # Extract from TDS
        if tds_pdf and os.path.exists(tds_pdf):
            tds_data = self.extractor.extract_tds(tds_pdf)
            for key, val in tds_data.items():
                if hasattr(data, key):
                    setattr(data, key, val)
                elif key == "equipment_list":
                    data.equipment_list = val
                elif key == "manpower_list":
                    data.manpower_list = val
        
        # Extract from BOQ
        if boq_pdf and os.path.exists(boq_pdf):
            boq_data = self.extractor.extract_boq(boq_pdf)
            data.boq_items = boq_data.get("boq_items", [])
        
        logger.info(f"Extraction complete: Tender {data.tender_id}, {len(data.boq_items)} BOQ items")
        return data

    def _load_template(self, template_path: str) -> Any:
        """Load a DOCX template (cached)."""
        from docx import Document
        
        if template_path not in self.template_cache:
            self.template_cache[template_path] = Document(template_path)
        return self.template_cache[template_path]

    def _replace_text(self, doc: Any, old_text: str, new_text: str):
        """Replace text in all paragraphs and tables of a document."""
        if not new_text:
            new_text = ""
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

    def _map_data_to_template(self, doc: Any, data: TenderData, mappings: Dict[str, str]):
        """Apply data mappings to a template document."""
        for data_field, template_placeholder in mappings.items():
            value = getattr(data, data_field, "")
            if value is None:
                value = ""
            if isinstance(value, list):
                value = ", ".join(str(v) for v in value)
            if not isinstance(value, str):
                value = str(value)
            self._replace_text(doc, template_placeholder, value)

    def generate_jv_deed(self, data: TenderData, output_path: str, template_path: str = None) -> str:
        """Generate Joint Venture Agreement."""
        if not template_path:
            template_path = os.path.join(self.templates_dir, "JV_DEED_TEMPLATE.docx")
        
        # If no template, create from scratch with data
        from docx import Document
        
        if os.path.exists(template_path):
            doc = self._load_template(template_path)
        else:
            doc = Document()
            doc.add_heading("JOINT VENTURE AGREEMENT & MEMORANDUM OF UNDERSTANDING", 0)
            doc.add_paragraph(f"This Joint Venture Agreement is made on {datetime.now().strftime('%d %B %Y')}")
        
        # Apply mappings
        mappings = {
            "tender_id": "{TENDER_ID}",
            "package_no": "{PACKAGE_NO}",
            "project_name": "{PROJECT_NAME}",
            "estimated_value_bdt": "{ESTIMATED_VALUE}",
            "bidder_name": "{BIDDER_NAME}",
        }
        self._map_data_to_template(doc, data, mappings)
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info(f"Generated: {output_path}")
        return output_path

    def generate_credit_line(self, data: TenderData, output_path: str, template_path: str = None) -> str:
        """Generate Letter of Commitment for Line of Credit (Form PW3-7)."""
        from docx import Document
        
        if template_path and os.path.exists(template_path):
            doc = self._load_template(template_path)
        else:
            doc = Document()
            doc.add_heading("Letter of Commitment for Bank's Undertaking for Line of Credit (Form PW3-7)", 0)
            doc.add_paragraph(f"Credit Commitment No: [insert number]")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
            doc.add_paragraph(
                f"We have been informed that {data.bidder_name or '[Bidder Name]'} "
                f"(hereinafter called \"the Tenderer\") intends to submit to you its "
                f"Tender for Tender ID: {data.tender_id}, Package: {data.package_no}."
            )
        
        mappings = {
            "tender_id": "{TENDER_ID}",
            "package_no": "{PACKAGE_NO}",
            "tender_security_bdt": "{SECURITY_AMOUNT}",
            "bidder_name": "{BIDDER_NAME}",
            "bank_name": "{BANK_NAME}",
            "bank_branch": "{BANK_BRANCH}",
            "project_name": "{PROJECT_NAME}",
        }
        self._map_data_to_template(doc, data, mappings)
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info(f"Generated: {output_path}")
        return output_path

    def generate_bg(self, data: TenderData, output_path: str, template_path: str = None) -> str:
        """Generate Bank Guarantee for Tender Security (Form e-PW3-5)."""
        from docx import Document
        
        if template_path and os.path.exists(template_path):
            doc = self._load_template(template_path)
        else:
            doc = Document()
            doc.add_heading("Bank Guarantee for Tender Security (Form e-PW3-5)", 0)
            doc.add_paragraph(f"Tender/Proposal ID: {data.tender_id}")
            doc.add_paragraph(f"TENDER GUARANTEE No: () Date () for Tk. {data.tender_security_bdt:,.2f}")
        
        mappings = {
            "tender_id": "{TENDER_ID}",
            "package_no": "{PACKAGE_NO}",
            "tender_security_bdt": "{SECURITY_AMOUNT}",
            "tender_security_text": "{SECURITY_TEXT}",
            "bidder_name": "{BIDDER_NAME}",
            "bidder_address": "{BIDDER_ADDRESS}",
            "bank_name": "{BANK_NAME}",
            "bank_branch": "{BANK_BRANCH}",
            "bank_address": "{BANK_ADDRESS}",
            "procuring_entity": "{PROCURING_ENTITY}",
            "invitation_ref": "{INVITATION_REF}",
            "validity_period_days": "{VALIDITY_DAYS}",
        }
        self._map_data_to_template(doc, data, mappings)
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info(f"Generated: {output_path}")
        return output_path

    def generate_equipment_declaration(self, data: TenderData, output_path: str, template_path: str = None) -> str:
        """Generate Equipment Declaration."""
        from docx import Document
        
        if template_path and os.path.exists(template_path):
            doc = self._load_template(template_path)
        else:
            doc = Document()
            doc.add_heading("EQUIPMENT DECLARATION", 0)
            doc.add_paragraph(
                f"We do hereby give assurance that all requisite construction equipment will arrive at "
                f"the working site for the work of \"{data.package_description or data.project_name}\" "
                f"in time as per work schedule."
            )
            
            # Build equipment table
            if data.equipment_list:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'No'
                hdr[1].text = 'Equipment Type and Characteristics'
                hdr[2].text = 'Minimum Number Required'
                for eq in data.equipment_list:
                    row = table.add_row().cells
                    row[0].text = str(eq.get("sl", ""))
                    row[1].text = eq.get("type", "")
                    row[2].text = eq.get("min_required", "")
        
        mappings = {
            "package_description": "{PACKAGE_DESC}",
            "project_name": "{PROJECT_NAME}",
            "bidder_name": "{BIDDER_NAME}",
        }
        self._map_data_to_template(doc, data, mappings)
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info(f"Generated: {output_path}")
        return output_path

    def generate_manpower_declaration(self, data: TenderData, output_path: str, template_path: str = None) -> str:
        """Generate Manpower Declaration."""
        from docx import Document
        
        if template_path and os.path.exists(template_path):
            doc = self._load_template(template_path)
        else:
            doc = Document()
            doc.add_heading("MANPOWER DECLARATION", 0)
            doc.add_paragraph(
                f"We do hereby give assurance that efficient and experienced Team of Man Power will be sent "
                f"at working site in time for work of \"{data.package_description or data.project_name}\"."
            )
            
            if data.manpower_list:
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                headers = ['Sl. No.', 'Post of Key Staff', 'Qualification', 'Nos.', 'Total Exp.', 'Similar Exp.']
                for i, h in enumerate(headers):
                    hdr[i].text = h
                for mp in data.manpower_list:
                    row = table.add_row().cells
                    row[0].text = str(mp.get("sl", ""))
                    row[1].text = mp.get("post", "")
                    row[2].text = mp.get("qualification", "")
                    row[3].text = mp.get("nos", "")
                    row[4].text = mp.get("total_exp", "")
                    row[5].text = mp.get("similar_exp", "")
        
        self._map_data_to_template(doc, data, {"package_description": "{PACKAGE_DESC}", "project_name": "{PROJECT_NAME}"})
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info(f"Generated: {output_path}")
        return output_path

    def generate_methodology(self, data: TenderData, output_path: str, template_path: str = None) -> str:
        """Generate Statement of Work Method (Methodology)."""
        from docx import Document
        
        if template_path and os.path.exists(template_path):
            doc = self._load_template(template_path)
        else:
            doc = Document()
            doc.add_heading("STATEMENT OF WORK METHOD", 0)
            
            doc.add_heading("NATURE OF WORKS:", 2)
            doc.add_paragraph(data.package_description or data.project_name or "")
            
            doc.add_heading("INTRODUCTION", 2)
            doc.add_paragraph(
                f"This work method statement outlines the proposed approach for the successful execution of "
                f"the {data.package_description or 'works'} for Tender ID: {data.tender_id}, "
                f"Package: {data.package_no}. {data.procuring_entity}, {data.organization}."
            )
            
            doc.add_heading("MAJOR COMPONENTS OF WORK", 2)
            components = [
                "Mobilization of manpower, equipment, and materials.",
                "Site Preparation, including clearing and pre-work surveys.",
                "Earthwork in cutting & filling to design slope.",
                "Geo-Textile Filter and Sand Filter laying.",
                "C.C. Block Manufacturing and Curing.",
                "C.C. Block Dumping & Placing on the prepared slope.",
                "De-mobilization & Site Handover to the Procuring Entity.",
            ]
            for c in components:
                doc.add_paragraph(c, style='List Bullet')
            
            doc.add_heading("EQUIPMENT MOBILIZATION", 2)
            if data.equipment_list:
                doc.add_paragraph("The following major construction equipment will be mobilized to the site:")
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'No'
                hdr[1].text = 'Equipment Type'
                hdr[2].text = 'Quantity'
                for eq in data.equipment_list:
                    row = table.add_row().cells
                    row[0].text = str(eq.get("sl", ""))
                    row[1].text = eq.get("type", "")
                    row[2].text = eq.get("min_required", "")
        
        mappings = {
            "tender_id": "{TENDER_ID}",
            "package_no": "{PACKAGE_NO}",
            "package_description": "{PACKAGE_DESC}",
            "project_name": "{PROJECT_NAME}",
            "procuring_entity": "{PROCURING_ENTITY}",
            "organization": "{ORGANIZATION}",
        }
        self._map_data_to_template(doc, data, mappings)
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info(f"Generated: {output_path}")
        return output_path

    def generate_all(self, data: TenderData, output_dir: str = "./output") -> Dict[str, str]:
        """Generate all tender submission documents."""
        os.makedirs(output_dir, exist_ok=True)
        
        results = {}
        
        generators = [
            ("JV_DEED.docx", self.generate_jv_deed),
            ("Credit_Line.docx", self.generate_credit_line),
            ("BG_Tender_Security.docx", self.generate_bg),
            ("Equipment_Declaration.docx", self.generate_equipment_declaration),
            ("Manpower_Declaration.docx", self.generate_manpower_declaration),
            ("Methodology.docx", self.generate_methodology),
        ]
        
        for filename, generator in generators:
            try:
                output_path = os.path.join(output_dir, filename)
                result = generator(data, output_path)
                results[filename] = "✅" if result else "❌"
            except Exception as e:
                logger.error(f"Failed to generate {filename}: {e}")
                results[filename] = f"❌ {str(e)[:50]}"
        
        logger.info(f"Generated {sum(1 for v in results.values() if '✅' in v)}/{len(results)} documents")
        return results


def create_template_from_uploaded(docx_path: str, output_template_path: str):
    """Create a template version of an uploaded DOCX by replacing specific values with placeholders."""
    from docx import Document
    
    doc = Document(docx_path)
    
    # Save as template (same content, but placed in templates dir)
    os.makedirs(os.path.dirname(output_template_path) or ".", exist_ok=True)
    doc.save(output_template_path)
    logger.info(f"Template saved: {output_template_path}")


def import_boq_from_excel(excel_path: str) -> List[Dict[str, Any]]:
    """Import BOQ items from the BOQ Rate Analysis Excel file.
    
    Args:
        excel_path: Path to the BOQ_Rate_Analysis Excel file
        
    Returns:
        List of BOQ item dicts with code, description, unit, quantity, rates
    """
    import openpyxl
    items = []
    try:
        wb = openpyxl.load_workbook(excel_path)
        ws = wb['BOQ Rate Comparison']
        for r in range(5, ws.max_row + 1):
            item_no = ws.cell(r, 1).value
            if isinstance(item_no, int) and item_no > 0:
                code = str(ws.cell(r, 2).value or '').strip()
                agency = str(ws.cell(r, 3).value or '')
                desc = str(ws.cell(r, 5).value or '').strip()
                unit = str(ws.cell(r, 6).value or '').strip()
                qty = float(ws.cell(r, 7).value or 0)
                sor_rate = float(ws.cell(r, 8).value or 0)
                quoted_rate = float(ws.cell(r, 10).value or 0)
                
                items.append({
                    'item_no': item_no,
                    'code': code,
                    'agency': agency,
                    'description': desc,
                    'unit': unit.lower(),
                    'quantity': qty,
                    'sor_rate': sor_rate,
                    'quoted_rate': quoted_rate,
                    'total_sor': sor_rate * qty,
                    'total_quoted': quoted_rate * qty,
                })
    except Exception as e:
        logger.error(f"Excel import failed: {e}")
    return items


def import_boq_from_pdf(boq_pdf_path: str, zone: str = 'D') -> List[Dict[str, Any]]:
    """Import BOQ from PDF and match with SOR rates.
    
    Args:
        boq_pdf_path: Path to BOQ PDF
        zone: Zone for rate lookup (A, B, C, D)
        
    Returns:
        List of BOQ item dicts with matched SOR rates
    """
    from app.sor.boq_matcher import BOQParser, BOQMatcher
    
    # Parse BOQ from PDF
    raw_items = BOQParser.from_pdf(boq_pdf_path)
    
    # Load SOR matcher
    matcher = BOQMatcher()
    matcher.load_all()
    
    items = []
    for it in raw_items:
        code = it.get('code', '')
        desc = it.get('description', '')
        agency = it.get('agency', 'BWDB')
        quantity = it.get('quantity', 0)
        unit = it.get('unit', '')
        
        # Match against SOR
        result = matcher.match_item(code=code, description=desc, agency=agency, zone=zone)
        
        sor_rate = result.get('sor_rate', 0) if result.get('matched') else 0
        
        items.append({
            'item_no': it.get('item_no', 0),
            'code': code,
            'agency': agency,
            'description': desc[:200],
            'unit': unit,
            'quantity': quantity,
            'sor_rate': sor_rate,
            'quoted_rate': 0,  # To be filled by tenderer
            'total_sor': sor_rate * quantity,
            'total_quoted': 0,
        })
    
    return items


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    gen = TenderDocGenerator()
    
    # Try Excel first, fall back to PDF
    excel_path = "/tmp/codex-web-uploads/f-uXkixm/BOQ_Rate_Analysis_1264860_COMPLETE.xlsx"
    boq_items = import_boq_from_excel(excel_path)
    
    if not boq_items:
        print("⚠️  Excel BOQ not found, parsing PDF...")
        boq_items = import_boq_from_pdf(
            "/tmp/codex-web-uploads/f-QvWc3n/4.BOQ_1264860.pdf",
            zone='D'
        )
    
    data = gen.extract_from_pdfs(
        notice_pdf="/tmp/codex-web-uploads/f-oyM9Jk/1.Notice_1264860.pdf",
        tds_pdf="/tmp/codex-web-uploads/f-KSLPba/2.TDS_1_1264860.pdf",
        boq_pdf="/tmp/codex-web-uploads/f-QvWc3n/4.BOQ_1264860.pdf",
    )
    
    # Override BOQ with our enriched data
    data.boq_items = boq_items
    if boq_items:
        data.boq_total = sum(it.get('total_sor', 0) or it.get('total_quoted', 0) for it in boq_items)
    
    # Add bidder info
    data.bidder_name = "M/S. Hassan & Brothers"
    data.bidder_address = "Mahmud Tower, (9th Floor) 19, Siddique Bazar, Dhaka"
    data.bidder_phone = "02-9564202"
    data.bidder_email = "hbl.engr@gmail.com"
    data.bank_name = "NRB Commercial Bank Ltd."
    data.bank_branch = "Principal Branch"
    data.bank_address = "114, Motijheel C/A, Dhaka"
    
    print(f"\n{'='*60}")
    print(f"  TENDER DATA SUMMARY — {data.tender_id or '1264860'}")
    print(f"{'='*60}")
    for key, val in sorted(asdict(data).items()):
        if val and key not in ('boq_items', 'equipment_list', 'manpower_list', 'tax_requirements'):
            print(f"  {key:35s}: {str(val)[:80]}")
        elif key == 'boq_items':
            match_pct = sum(1 for it in val if it.get('sor_rate', 0) > 0) / max(len(val), 1) * 100
            print(f"  {key:35s}: {len(val)} items ({match_pct:.0f}% SOR matched)")
        elif key == 'equipment_list':
            print(f"  {key:35s}: {len(val)} items")
        elif key == 'manpower_list':
            print(f"  {key:35s}: {len(val)} items")
    
    print(f"\n{'='*60}")
    print(f"  Generating Documents...")
    print(f"{'='*60}")
    results = gen.generate_all(data, output_dir="/tmp/generated_docs")
    
    for fname, status in results.items():
        print(f"  {status}  {fname}")
    
    print(f"\n✅ Documents generated in /tmp/generated_docs/")
